import sys
import subprocess

# --- AUTO-INSTALLER: Forces the current Python brain to get the library ---
try:
    import docx
except ModuleNotFoundError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

import os
import streamlit as st
import pandas as pd
from dotenv import load_dotenv
import google.generativeai as genai
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import json
from io import BytesIO

# This must be the very first Streamlit command!
st.set_page_config(page_title="Dremel AI Hub", page_icon="🔧", layout="wide")

# --- 1. CONFIGURATION & AI ---
load_dotenv()

# Try local .env first. If it fails (because we are on the cloud), use Streamlit Secrets.
GEMINI_KEY = os.getenv('GEMINI_API_KEY')
if not GEMINI_KEY:
    try:
        GEMINI_KEY = st.secrets["GEMINI_API_KEY"]
    except Exception:
        pass

if GEMINI_KEY:
    genai.configure(api_key=GEMINI_KEY)

# --- 2. DREMEL THEME INJECTION ---
st.markdown("""
    <style>
        /* 1. Dremel Cyan Accent for main buttons */
        div.stButton > button:first-child {
            background-color: #00AEEF !important;
            color: white !important;
            border: none !important;
            font-weight: bold;
            padding: 10px 20px;
            border-radius: 5px;
        }
        div.stButton > button:first-child:hover {
            background-color: #0089bd !important;
        }

        /* 2. Main page Cyan separators */
        hr {
            border-top: 2px solid #00AEEF;
        }

        /* 3. DREMEL BLUE SIDEBAR */
        [data-testid="stSidebar"] {
            background-color: #00529B !important;
        }

        /* 4. Force all text in the sidebar to be crisp white */
        [data-testid="stSidebar"] p, 
        [data-testid="stSidebar"] h1, 
        [data-testid="stSidebar"] h2, 
        [data-testid="stSidebar"] h3,
        [data-testid="stSidebar"] label {
            color: #FFFFFF !important;
        }

        /* 5. Make the sidebar dividers Cyan */
        [data-testid="stSidebar"] hr {
            border-top: 2px solid #00AEEF !important;
        }

        /* Tighten up the top margin */
        .block-container {
            padding-top: 2rem !important;
        }
    </style>
""", unsafe_allow_html=True)


# --- 3. CLOUD DATA CONNECTION ---
@st.cache_data(ttl=3600)
def load_cloud_data():
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]

    if os.path.exists("dremel-bot-key.json"):
        creds = ServiceAccountCredentials.from_json_keyfile_name("dremel-bot-key.json", scope)
    else:
        creds_dict = json.loads(st.secrets["google_sheets"]["json_key"])
        creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)

    client = gspread.authorize(creds)
    sheet = client.open("Dremel_Trending_Data").sheet1
    data = sheet.get_all_records()
    return pd.DataFrame(data)


# --- 4. BRANDING & MAIN NAVIGATION ---
try:
    st.sidebar.image("dremel_logo.png", use_container_width=True)
except Exception:
    st.sidebar.markdown("## DREMEL")

st.sidebar.markdown("---")
st.sidebar.markdown("### 🎛️ App Navigation")

# REPLACED SELECTBOX WITH RADIO BUTTONS TO REMOVE SEARCH CAPABILITY
page_selection = st.sidebar.radio(
    "Select a Page:",
    [
        "🚀 Live Tools & About",
        "🏗️ Architecture: Dashboard",
        "🏗️ Architecture: AI Engine"
    ]
)

st.sidebar.markdown("---")
st.sidebar.info("Developed by Mihir Bose, Powered by AI")

# ==========================================
# PAGE 1: THE MAIN LIVE TOOLS & ABOUT TABS
# ==========================================
if page_selection == "🚀 Live Tools & About":

    # THE NEW MASTER TABS AT THE TOP OF THE PAGE
    main_tab1, main_tab2 = st.tabs(["🚀 Live Dashboard & AI Engine", "ℹ️ About the Project"])

    # ------------------------------------------
    # SUB-TAB 1: LIVE TOOLS
    # ------------------------------------------
    with main_tab1:
        st.sidebar.markdown("**Tool Display Toggles:**")
        show_dashboard = st.sidebar.toggle("📊 Live Dashboard", value=True)
        show_ai = st.sidebar.toggle("🤖 AI Content Generator", value=True)

        if show_dashboard:
            st.markdown("""
                <div style="background-color: #00529B; padding: 20px; border-radius: 8px; margin-bottom: 25px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
                    <h2 style="color: #FFFFFF; margin: 0;">📊 Live Trend Engine Dashboard</h2>
                    <p style="color: #FFFFFF; margin: 0; font-size: 16px;">Interact with the live data below to spot the fastest-growing DIY categories.</p>
                </div>
            """, unsafe_allow_html=True)

            TABLEAU_URL = "https://public.tableau.com/views/Dremel_TrendsDashboard/LiveDashboard?:language=en-US&publish=yes&:sid=&:redirect=auth&:display_count=n&:origin=viz_share_link?:embed=yes&:showVizHome=no"
            st.components.v1.iframe(TABLEAU_URL, width=1200, height=700, scrolling=True)

        if show_dashboard and show_ai:
            st.markdown("<br><br>", unsafe_allow_html=True)
            st.markdown("---")
            st.markdown("<br><br>", unsafe_allow_html=True)

        if show_ai:
            st.markdown("""
                <div style="background-color: #00529B; padding: 20px; border-radius: 8px; margin-bottom: 25px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
                    <h2 style="color: #FFFFFF; margin: 0;">🤖 Dremel AI Ideation Center 🚀</h2>
                    <p style="color: #FFFFFF; margin: 0; font-size: 16px;">Instantly translate rising trends into cross-channel marketing strategies.</p>
                </div>
            """, unsafe_allow_html=True)

            try:
                df = load_cloud_data()
                unique_keywords = sorted(df['Keyword'].dropna().unique())

                st.subheader("🎯 Step 1: Target a Trend")

                # --- REVERTED BACK TO A SINGLE SEARCH/DROPDOWN BOX ---
                selected_trend = st.selectbox(
                    "Click inside the box to select a trend from the list, OR type directly to search for one:",
                    unique_keywords
                )

                st.write("")

                if "ai_raw_text" not in st.session_state:
                    st.session_state.ai_raw_text = None
                if "current_trend" not in st.session_state:
                    st.session_state.current_trend = None

                # Reset the text if the user picks a new trend
                if st.session_state.current_trend != selected_trend:
                    st.session_state.ai_raw_text = None
                    st.session_state.current_trend = selected_trend

                if st.button("⚡ Activate AI Virtual Managers", type="primary"):
                    with st.spinner(f"Coordinating with virtual managers to analyze '{selected_trend}'..."):
                        prompt = f"""
                        You are a team of expert digital marketing managers at Dremel UK. 
                        The automated data engine has identified '{selected_trend}' as a major rising trend on YouTube UK.
                        Provide actionable, specific tasks for each of the following 6 managers. Keep each role's response to 2-3 clear, high-impact bullet points and use professional language tailored to Dremel's rotary tools and solutions.
                        Format your entire response EXACTLY like this so the app can read it:
                        [CONTENT]
                        - Bullet 1
                        [SOCIAL]
                        - Bullet 1
                        [EMAIL]
                        - Bullet 1
                        [WEBSITE]
                        - Bullet 1
                        [ECOMMERCE]
                        - Bullet 1
                        [PRODUCT]
                        - Bullet 1
                        """
                        model = genai.GenerativeModel('gemini-2.5-flash')
                        response = model.generate_content(prompt)
                        st.session_state.ai_raw_text = response.text

                # --- RENDER THE TABS & DOCUMENT GENERATOR ---
                if st.session_state.ai_raw_text:
                    def get_section(text, tag):
                        try:
                            parts = text.split(f"[{tag}]")
                            if len(parts) > 1:
                                return parts[1].split("[")[0].strip()
                            return "Strategy formulation in progress..."
                        except:
                            return "Strategy formulation in progress..."


                    content_strat = get_section(st.session_state.ai_raw_text, "CONTENT")
                    social_strat = get_section(st.session_state.ai_raw_text, "SOCIAL")
                    email_strat = get_section(st.session_state.ai_raw_text, "EMAIL")
                    web_strat = get_section(st.session_state.ai_raw_text, "WEBSITE")
                    ecom_strat = get_section(st.session_state.ai_raw_text, "ECOMMERCE")
                    prod_strat = get_section(st.session_state.ai_raw_text, "PRODUCT")

                    st.markdown("---")
                    st.subheader(f"📋 Operational Action Plan: {selected_trend.upper()}")

                    tabs = st.tabs([
                        "Content Manager", "Social Media", "Email Marketing",
                        "Website Manager", "3rd Party E-Com", "Product Manager"
                    ])

                    manager_data = [
                        ("Content Manager", "🎬 Video & Content Strategy", content_strat),
                        ("Social Media", "📱 Social Media & Influencer Partnerships", social_strat),
                        ("Email Marketing", "✉️ Email Campaign & Retention", email_strat),
                        ("Website Manager", "🌐 Homepage & Landing Page Optimization", web_strat),
                        ("3rd Party E-Com", "🛍️ Amazon & 3rd Party Retail Channels", ecom_strat),
                        ("Product Manager", "🛠️ Product Bundling & Accessory Focus", prod_strat)
                    ]


                    def create_docx(manager_role, strategy_bullets, trend):
                        deep_dive_prompt = f"""
                        You are a Senior {manager_role} at Dremel. You need to hand off the following strategy to a junior coordinator for immediate execution.
                        The trending topic is: {trend}.
                        The high-level strategy is: 
                        {strategy_bullets}

                        TONE & SAFETY DIRECTIVE:
                        Analyze the trend '{trend}'. 
                        - If it involves heavy-duty, complex, or serious DIY work (e.g., cutting, metalwork, routing), adopt a highly professional tone and include strict, specific safety precautions (PPE, proper tool handling).
                        - If it is a fun, craft-based, or light make-at-home project (e.g., jewelry, simple decor), adopt an enthusiastic, creative tone, but still explicitly mention essential basic safety tips.

                        Write a highly detailed, step-by-step implementation brief. The structure must be clear and leave zero confusion for the junior employee.
                        Include exactly these sections:
                        1. Exact Actionable Steps for each bullet point.
                        2. Specifications: Required channels, image aspect ratios, hashtags, specific Dremel tool requirements, or design guidelines where relevant.
                        3. UGC Strategy: A specific playbook on how to source, incentivize, and integrate User Generated Content (UGC) for this campaign.
                        4. Success Metrics: KPIs the junior employee should track.

                        Format it cleanly with standard spacing. Do not use markdown symbols like asterisks or hashtags in the text body.
                        """
                        model = genai.GenerativeModel('gemini-2.5-flash')
                        detailed_response = model.generate_content(deep_dive_prompt)

                        doc = docx.Document()
                        doc.add_heading(f'Execution Brief: {manager_role}', 0)
                        doc.add_heading(f'Campaign Trend: {trend}', 1)
                        doc.add_paragraph(detailed_response.text)

                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        return buffer


                    for i, (role, title, strat) in enumerate(manager_data):
                        with tabs[i]:
                            st.info(title)
                            st.write(strat)
                            st.write("")

                            if st.button(f"📥 Draft Detailed .docx Brief for {role}", key=f"btn_{i}"):
                                with st.spinner(
                                        "Analyzing keyword tone and generating junior-level implementation brief..."):
                                    docx_file = create_docx(role, strat, selected_trend)
                                    st.download_button(
                                        label=f"✅ Click to Download {role} Brief",
                                        data=docx_file,
                                        file_name=f"Dremel_{role.replace(' ', '_')}_Brief.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )

            except Exception as e:
                st.error(f"Error connecting to Cloud Data: {e}")

    # ------------------------------------------
    # SUB-TAB 2: ABOUT THE PROJECT
    # ------------------------------------------
    with main_tab2:
        st.markdown("""
            <div style="background-color: #00529B; padding: 20px; border-radius: 8px; margin-bottom: 25px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
                <h2 style="color: #FFFFFF; margin: 0;">ℹ️ About This Project</h2>
                <p style="color: #FFFFFF; margin: 0; font-size: 16px;">The context, architecture, and optimal usage of the Dremel AI Hub.</p>
            </div>
        """, unsafe_allow_html=True)

        st.markdown("### 1. Why this platform was created?")
        st.write("""
        This platform was developed to modernize Dremel's marketing strategy by leveraging data-driven insights and artificial intelligence. 
        Instead of relying on historical assumptions or manual research, this tool actively monitors live consumer interests and DIY trends in the UK. 
        The goal is to bridge the gap between what target audiences are actively searching for and the multi-channel campaigns Dremel creates, ensuring maximum relevance, engagement, and ROI.
        """)

        st.divider()

        st.markdown("### 2. How it was created & how it works")
        st.write("""
        This application operates on a fully autonomous, production-grade cloud architecture built completely from scratch:
        * **Data Collection:** A serverless Python script runs nightly via GitHub Actions, scraping the YouTube API for the latest DIY and woodworking trends.
        * **NLP Processing:** The system uses Spacy (Natural Language Processing) to extract core topics and filter out conversational noise.
        * **Cloud Database:** The refined keyword data is autonomously pushed to a secure Google Sheets database.
        * **Data Visualization:** Tableau Public connects directly to this live database, rendering an interactive dashboard that updates automatically.
        * **AI Generation:** The frontend interfaces with Google's Gemini AI to dynamically generate tailored, cross-channel marketing strategies based on the live data, and now includes automated deep-dive Execution Briefs (.docx) for team hand-offs.
        """)
        st.info(
            "💡 *Tip: Check out the 'Architecture' pages in the sidebar menu for a deep dive into how the Dashboard and AI Engine code works!*")

        st.divider()

        st.markdown("### 3. How to use this website efficiently")
        st.markdown("""
        1. **Analyze the Data:** Go to the 'Live Tools' page and start by exploring the Tableau dashboard. Look for sudden spikes in specific categories to identify what is currently capturing audience attention.
        2. **Select a Trend:** Scroll down to the AI Ideation engine and select a high-value keyword from the dropdown (or type directly in the box to search).
        3. **Generate Strategy:** Click the **'⚡ Activate AI Virtual Managers'** button. The AI will output a customized, multi-channel campaign.
        4. **Execute:** Inside any manager's tab, click the download button to automatically draft a highly detailed, tone-appropriate implementation brief for a junior employee to execute immediately.
        """)


# ==========================================
# PAGE 2: DASHBOARD DOCUMENTATION
# ==========================================
elif page_selection == "🏗️ Architecture: Dashboard":
    st.markdown("""
        <div style="background-color: #00529B; padding: 20px; border-radius: 8px; margin-bottom: 25px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
            <h2 style="color: #FFFFFF; margin: 0;">🏗️ Data Architecture: The Trend Engine</h2>
            <p style="color: #FFFFFF; margin: 0; font-size: 16px;">Behind the scenes of the automated Tableau Dashboard.</p>
        </div>
    """, unsafe_allow_html=True)

    st.markdown("### 1. The Python Scraper (Data Collection)")
    st.write(
        "The system begins with a scheduled Python script running in the background. It utilizes the **Google API Client** to connect to YouTube's v3 API. The script is programmed with 25 highly specific DIY target keywords (e.g., 'Woodworking', 'Furniture Flips', 'Kids Crafts'). It queries the UK server region specifically, scraping the top 50 videos for each category to ensure a highly localized dataset.")

    st.markdown("### 2. Natural Language Processing (Data Processing)")
    st.write(
        "Raw video titles and descriptions are messy. To find true trends, the script uses **Spacy**, an advanced Natural Language Processing (NLP) library. The NLP engine strips away filler words ('stop words'), analyzes the grammar, and extracts only the core Nouns and Proper Nouns (e.g., pulling 'Epoxy' out of a long video title).")

    st.markdown("### 3. Google Cloud Integration (Data Storage)")
    st.write(
        "Once the data is cleaned into a Pandas DataFrame, the script authenticates securely with Google Cloud using a Service Account JSON Key and the `gspread` library. It acts as an automated 'robot worker,' wiping yesterday's data and writing today's fresh metrics directly to a hosted Google Sheet.")

    st.markdown("### 4. Tableau Public (Data Visualization)")
    st.write(
        "The final UI layer is built in Tableau. Because the dashboard is linked directly to the Google Sheet via an active data connection, it auto-syncs on a 24-hour cycle. The interactive iframe embedded in this app simply requests the most recent, published version of the dashboard from Tableau's servers.")

# ==========================================
# PAGE 3: AI ENGINE DOCUMENTATION
# ==========================================
elif page_selection == "🏗️ Architecture: AI Engine":
    st.markdown("""
        <div style="background-color: #00529B; padding: 20px; border-radius: 8px; margin-bottom: 25px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
            <h2 style="color: #FFFFFF; margin: 0;">🏗️ System Architecture: The AI Ideation Tool</h2>
            <p style="color: #FFFFFF; margin: 0; font-size: 16px;">How Large Language Models (LLMs) are used to automate marketing strategy.</p>
        </div>
    """, unsafe_allow_html=True)

    st.markdown("### 1. Dynamic Cloud Retrieval")
    st.write(
        "To ensure the AI is only generating strategies based on *actual* current trends, the app must read the database first. Using Streamlit's `@st.cache_data` command, the app connects to the same Google Sheet updated by the Tableau dashboard. It caches this data so the website loads instantly, then parses the columns to populate the user's dropdown search box with live keywords.")

    st.markdown("### 2. Prompt Engineering & API Connection")
    st.write(
        "When a user clicks 'Activate Virtual Managers', the system takes the selected keyword and injects it into a highly engineered system prompt. It connects to the **Google Gemini 3.5 Flash** model via the `google.generativeai` library. The prompt specifically instructs the AI to adopt the persona of a 'Dremel UK Digital Marketing Manager' and forces it to output its response in a very strict, machine-readable format (using tags like `[CONTENT]` and `[SOCIAL]`).")

    st.markdown("### 3. Data Parsing & Routing")
    st.write(
        "Large Language Models naturally output raw blocks of text. To create the 'Manager Tabs' UI, the Python backend intercepts the raw text from Google Gemini. It runs a custom `get_section()` parsing function to split the text block apart based on the hidden tags we forced the AI to use. Finally, it routes each specific section of text into Streamlit's isolated UI tabs, creating the illusion of 6 different virtual employees working simultaneously.")

    st.markdown("### 4. Automated Execution Briefs")
    st.write(
        "To close the gap between high-level strategy and immediate execution, the system features a document generation engine utilizing `python-docx` and `BytesIO`. When triggered, a secondary AI prompt acts as a Senior Manager, expanding the strategy into granular, step-by-step instructions. The AI analyzes the keyword to dynamically adjust its tone (e.g., professional safety warnings for heavy DIY vs. enthusiastic tone for crafts). The app builds a complete Microsoft Word `.docx` file in the server's short-term memory and securely delivers it directly to the user's desktop, leaving no physical files behind on the server.")
